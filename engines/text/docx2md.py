#!/usr/bin/env python3
"""
docx2md — Word (.docx) to Markdown. Sibling to pdf2md.py; shares the same
CLI contract (-o/--output, --quiet, exit codes) via common.py, so
pdf2md-auto.sh and any caller invoke it identically to the PDF engines
regardless of source format.

Unlike PDF, .docx is already a structured, text-native format -- no OCR, no
scanning, no rotation, and therefore no possibility of the OCR/LLM spelling-
decision defect this project exists to prevent (there's no recognition step
to get wrong; every character is read directly from the document XML).

The engineering concern here is COMPLETENESS, not recognition fidelity:
  - python-docx's document.paragraphs and document.tables are two SEPARATE
    lists that lose the true interleaving of narrative and tables -- read
    directly from the body XML in document order instead, or a financial
    statement's tables and surrounding text end up silently reordered
    relative to each other.
  - headers and footers are outside document.paragraphs entirely and are
    silently missed by a naive walk -- included here explicitly (and
    deduplicated, since Word repeats the same header/footer per section).
  - merged table cells are represented as the SAME cell object at every
    spanned position (docx's own structure, not a bug) -- reproduced as-is
    (duplicated across spanned columns) rather than blanked, since
    correctly detecting "is this cell part of a merge" is a heuristic that
    could misfire, and duplication is data-safe (never loses a value)
    where blanking on a wrong guess would not be.

Output is the same GFM pipe-table markdown syntax pdf2md's text engine
already produces, so downstream table-parsing code needs no changes to
handle docx-derived documents.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from common import run


def cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip().replace("|", "\\|")


def table_to_markdown(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    header = [cell_text(c) for c in rows[0].cells]
    if not header:
        return ""
    lines = ["| " + " | ".join(header) + " |",
             "|" + "|".join("---" for _ in header) + "|"]
    for row in rows[1:]:
        cells = [cell_text(c) for c in row.cells]
        # defensive padding -- a malformed row shouldn't crash the whole document
        if len(cells) < len(header):
            cells += [""] * (len(header) - len(cells))
        elif len(cells) > len(header):
            cells = cells[:len(header)]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def paragraph_to_markdown(p) -> str:
    text = p.text.strip()
    if not text:
        return ""
    style = (p.style.name or "").lower() if p.style else ""
    if style == "title" or "heading 1" in style:
        return f"# {text}"
    if "heading 2" in style:
        return f"## {text}"
    if "heading 3" in style:
        return f"### {text}"
    if style.startswith("heading"):
        return f"#### {text}"
    return text


def block_items_in_order(document):
    """Yield paragraphs and tables in TRUE document order by reading the
    body XML directly -- see module docstring."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def section_headers_footers(document) -> list[str]:
    """Headers/footers, deduplicated (Word repeats the same one per section
    unless explicitly varied) -- included, not dropped, but kept clearly
    separate from the main body so routine boilerplate (page numbers, doc
    title) doesn't get mistaken for narrative financial-statement text."""
    seen, blocks = set(), []
    for section in document.sections:
        for part, label in ((section.header, "Header"), (section.footer, "Footer")):
            if part.is_linked_to_previous:
                continue
            lines = [paragraph_to_markdown(p) for p in part.paragraphs]
            text = "\n".join(line for line in lines if line).strip()
            if text and text not in seen:
                seen.add(text)
                blocks.append(f"*{label}: {text}*")
    return blocks


def to_markdown(input_path: str) -> str:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(input_path)

    out = []
    hf = section_headers_footers(document)
    if hf:
        out.append("\n\n".join(hf))
        out.append("---")

    for item in block_items_in_order(document):
        if isinstance(item, Paragraph):
            md = paragraph_to_markdown(item)
            if md:
                out.append(md)
        elif isinstance(item, Table):
            md = table_to_markdown(item)
            if md:
                out.append(md)

    text = "\n\n".join(out)
    if len(text.strip()) < 20:
        raise ValueError(f"produced almost no output ({len(text.strip())} chars) -- "
                          f"this .docx may be empty, corrupted, or in a form python-docx can't read")
    return text


def _convert(args, log):
    log(f"[docx2md] reading {args.input}")
    out = to_markdown(args.input)
    log(f"[docx2md] extracted {len(out)} chars")
    return out


if __name__ == "__main__":
    run("docx2md", "Word (.docx) to Markdown.", _convert)
